#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import sys

def change_font_to_heiti(docx_path, output_path):
    """将Word文档中的所有字体改为黑体"""
    try:
        doc = Document(docx_path)
        
        # 修改段落字体
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                # 设置中文字体为黑体
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                # 设置英文字体
                run.font.name = 'SimHei'
        
        # 修改表格中的字体
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = '黑体'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                            run.font.name = 'SimHei'
        
        # 保存修改后的文档
        doc.save(output_path)
        print(f"成功将字体改为黑体：{output_path}")
        return True
    except Exception as e:
        print(f"Error: {str(e)}")
        return False

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python3 change_font_to_heiti.py <input.docx> <output.docx>")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    change_font_to_heiti(input_file, output_file)
