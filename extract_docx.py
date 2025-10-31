#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import sys
from docx import Document

def extract_text_from_docx(docx_path):
    """?docx?????????"""
    try:
        doc = Document(docx_path)
        
        # ??????
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():  # ???????
                full_text.append(para.text)
        
        # ??????
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    full_text.append(' | '.join(row_text))
        
        return '\n'.join(full_text)
    except Exception as e:
        return f"Error reading {docx_path}: {str(e)}"

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 extract_docx.py <docx_file>")
        sys.exit(1)
    
    docx_file = sys.argv[1]
    content = extract_text_from_docx(docx_file)
    print(content)
